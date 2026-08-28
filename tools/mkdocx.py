#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""Build the social kit: media/RESENTMENT-2.0.0-social-kit.docx

Ready-to-post copy for each platform, the asset guide, the talking points and
the honest limitations - in one document somebody can open, read and copy out
of without having to understand the codebase first.

Generated rather than hand-written for the same reason the site and the preview
images are: the numbers in it come from the project and go stale the moment a
test is added. Change them here, once.

  python tools/mkdocx.py
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as exc:
    print("cannot import python-docx: %s" % exc)
    print("if it is not installed:  python -m pip install python-docx")
    sys.exit(2)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "media")

VERSION = "2.0.0"
REPO    = "https://github.com/ni-sh-a-char/RESENTMENT---kernel"
DOCS    = "https://ni-sh-a-char.github.io/RESENTMENT---kernel/"

INK    = RGBColor(0x14, 0x18, 0x20)
MUTED  = RGBColor(0x55, 0x5D, 0x6B)
BRASS  = RGBColor(0x9A, 0x6A, 0x10)
TEAL   = RGBColor(0x0D, 0x6E, 0x6B)
RULE   = RGBColor(0xD8, 0xDC, 0xE4)

BODY = "Segoe UI"
MONO = "Consolas"


# ---------------------------------------------------------------- helpers

def shade(cell, hexcolour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def rule(doc, colour=RULE):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (colour[0], colour[1], colour[2]))
    bdr.append(bottom)
    pPr.append(bdr)
    p.paragraph_format.space_after = Pt(10)
    return p


def para(doc, text="", size=10.5, bold=False, colour=INK, font=BODY,
         space_after=8, space_before=0, italic=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = colour
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    return p


def h1(doc, text):
    rule(doc)
    return para(doc, text, size=19, bold=True, space_after=10, space_before=2)


def h2(doc, text):
    return para(doc, text, size=13.5, bold=True, colour=BRASS,
                space_before=14, space_after=6)


def h3(doc, text):
    return para(doc, text, size=11, bold=True, colour=TEAL,
                space_before=10, space_after=4)


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = BODY
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(3)
    return p


def block(doc, text, mono=True, fill="F4F6FA"):
    """A copy-out-of-me box. Everything a person is meant to paste into a
    social network lives in one of these, so it is obvious what is commentary
    and what is the post."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    c.text = ""
    first = True
    for line in text.split("\n"):
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        r = p.add_run(line)
        r.font.size = Pt(9.5 if mono else 10)
        r.font.name = MONO if mono else BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), MONO if mono else BODY)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def kv_table(doc, rows, widths=(2.1, 4.4)):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in rows:
        cells = t.add_row().cells
        for i, (cell, text) in enumerate(zip(cells, (k, v))):
            cell.width = Inches(widths[i])
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            r.font.name = BODY
            r.font.bold = (i == 0)
            r.font.color.rgb = INK if i else MUTED
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ------------------------------------------------------------------- copy

X_MAIN = """I rebuilt my hobby kernel into something I'd actually put an OS on.

RESENTMENT 2.0.0 — a capability-secure, AI-native kernel.

Three things a normal kernel can't say:

• Authority expires. Every capability carries a cryptographic seal with the
  deadline inside its MAC. Forget to revoke one and it stops working anyway.

• The whole machine has one SHA-256 root digest. Two machines in the same
  state hash identically — so system state is something you diff, not
  something you infer from logs.

• Inference is a scheduling class, not a userspace afterthought.

Boots on x86_64, ARM64 and RISC-V. SMP on all three. Loads a GGUF model and
runs a real transformer forward pass through its own operators, on a thread the
deadline admission controller accepted.

1440 host assertions + 180 driven through its own shell under QEMU.
Zero dependencies.

github.com/ni-sh-a-char/RESENTMENT---kernel"""

X_THREAD = """1/  Two years ago this repo printed one line of text at boot. That's it.
    One line, in VGA text mode.

    I just tagged v2.0.0. Here's what changed and what it cost. 🧵

2/  The idea: what if authority *expired by default*?

    Every capability in RESENTMENT carries a Kaalka seal with a validity
    window baked into its MAC. Not a revocation list somebody has to
    remember to sweep — the permission simply stops working.

    Widen the window by editing the struct and the seal is invalid.

3/  Second idea: the whole machine is one hash.

    Every kernel object is a Merkle node. The canonical encoding excludes
    timestamps, pointers and IDs — deliberately — so two machines that
    booted hours apart but reached the same configuration hash identically.

    `.digest` gives you the root. That's your whole system state.

4/  Third: inference gets its own scheduling class.

    A token loop isn't interactive (it never sleeps for a human) and it
    isn't batch (it has a rate you can see). It's a soft real-time stream.

    SCHED_INFERENCE admits it with a declared rate, gives it a budget, and
    demotes rather than drops it when it overruns.

5/  Three architectures: x86_64, ARM64, RISC-V. All boot to a shell, all
    use every core.

    x86_64: ACPI MADT + a real-mode AP trampoline
    ARM64:  PSCI CPU_ON
    RISC-V: SBI HSM hart_start

6/  The bugs were the interesting part.

    The scheduler returned a thread to the run queue *before* its stack
    pointer was saved. Two cores could end up running one stack. It showed
    up as a page fault on cpu1, minutes later, in unrelated code.

7/  The x86_64 syscall entry destroyed rdi, rsi, rdx, r8, r9, r10.

    It built its argument block by pushing them and dropped it with
    `add rsp, 7*8` instead of popping them back.

    Symptom: a program printing the same wrong number five times while the
    string literals beside it came out perfectly.

8/  My favourite: the CSPRNG was running on an all-zero key.

    Entropy went into a pool. The pool was only folded into the key once
    the entropy estimate crossed 128 bits — which never happens without a
    hardware RNG. The kernel warned "entropy pool is weak."

    It was much worse than weak.

9/  My actual favourite, found while making inference work:

    sched_tick() was never called on ARM64. At all. Since the port was
    written.

    It was driven by a hardcoded test for interrupt line zero — the x86
    timer and nothing else's. ARM's timer is a PPI on line 27.

    So that port had no preemption, no slice accounting, and
    sched_sleep_ms never returned, because nothing ever woke a sleeper.

    Nothing in the test suite slept. It went unnoticed for the life of
    the port.

10/ And a RISC-V one that only appears under GNU ld:

    `la t0, sym` gets relaxed into `addi t0, gp, offset`. I'd put code
    before gp was initialised. Every hart computed a garbage address and
    parked itself.

    Invisible under LLD. Fatal on real hardware.

11/ The kernel now runs a transformer.

    Not a wrapper around one — the forward pass itself, built from the
    kernel's own operators. Embed, RMSNorm, QKV, RoPE, attention over the
    KV history, gated FFN, logits, argmax.

      resentment> .infer 24
        generated  24 tokens in 130 ms
        class      SCHED_INFERENCE, 50 Hz declared, admitted
        deadline miss  0

    "admitted" is the deadline admission controller accepting the rate.
    Ask for more than the class has left and it refuses.

12/ What I'm NOT claiming:

    The model has pseudo-random weights. The tokens mean nothing. The
    claim is that the *path* exists and is scheduled correctly — a
    systems claim, which is the only kind a kernel gets to make.

    No PCI, no network stack yet.

13/ Everything is verifiable:

    make test           1440 assertions on the host
    make qemu-test-all  6 targets, 180 assertions through the shell
    make kaalka-check   crypto byte-identical to the reference

    Apache 2.0. No dependencies. One command builds all three arches.

    github.com/ni-sh-a-char/RESENTMENT---kernel"""

LINKEDIN = """Two years ago I wrote a hobby kernel that printed one line of text at boot.

Last night I tagged v2.0.0, and it's a different thing entirely.

RESENTMENT is a capability-secure, AI-native kernel. It boots on x86_64, ARM64
and RISC-V, uses every core on the machine, and is verified by 1,440 assertions
on the host plus 180 driven through its own shell under QEMU on six targets.

Three design decisions it's built around:

𝗔𝘂𝘁𝗵𝗼𝗿𝗶𝘁𝘆 𝗲𝘅𝗽𝗶𝗿𝗲𝘀. There's no ambient authority anywhere in the system. A task
can't open a file because it knows the path — it can only act on objects it
holds a capability for, and every capability carries a cryptographic seal with
its deadline inside the MAC. A permission you forgot about stops working on its
own. Revocation as a background property rather than a chore.

𝗧𝗵𝗲 𝘀𝘆𝘀𝘁𝗲𝗺 𝗶𝘀 𝗮 𝗴𝗿𝗮𝗽𝗵 𝘆𝗼𝘂 𝗰𝗮𝗻 𝗵𝗮𝘀𝗵. Every kernel object is a Merkle node. The
canonical encoding deliberately excludes timestamps, pointers and IDs, so two
machines that booted at different times but reached the same configuration
produce the same root digest. That single property turns system state into
something you can diff, attest and replay instead of something you infer from
logs.

𝗜𝗻𝗳𝗲𝗿𝗲𝗻𝗰𝗲 𝗶𝘀 𝗮 𝘀𝗰𝗵𝗲𝗱𝘂𝗹𝗶𝗻𝗴 𝗰𝗹𝗮𝘀𝘀. A token-generation loop is neither
interactive nor batch — it's a soft real-time stream with a rate the user can
see. SCHED_INFERENCE treats it as one: admitted with a declared rate, given a
per-period budget, demoted rather than dropped when it overruns.

It also runs a transformer - not a wrapper around one, but the forward pass
itself, assembled from the kernel's own operators and scheduled by the class
above. The fixture model it runs has pseudo-random weights and produces
meaningless tokens, deliberately: the claim is that the path exists and is
admitted, budgeted and yielded correctly, which is a systems claim and the only
kind a kernel is entitled to make.

What I found more valuable than any feature was what building it surfaced. The
scheduler was returning a thread to the run queue before its stack pointer had
been saved — two cores could run one stack. The x86_64 system call entry was
destroying the argument registers a caller is entitled to keep live. The
CSPRNG was running on an all-zero key, because entropy only reached the key
once an estimate crossed a threshold that a machine without a hardware RNG
never reaches. And on ARM64 the scheduler tick was never called at all - it was
driven by a hardcoded test for interrupt line zero, which is the x86 timer and
nothing else's - so that port had no preemption and no way to wake a sleeping
thread, for its entire life, unnoticed because nothing in the suite slept.

Each of those passed review, passed tests, and was found only by a second
toolchain, a four-core workload, or by building something that finally
exercised the path.

I'm equally clear about what isn't done. There's no PCI enumeration and no network stack. The roadmap
lists those, and also lists what the project has decided not to build and why.

Apache 2.0, no dependencies, one command builds all three architectures.

Code: github.com/ni-sh-a-char/RESENTMENT---kernel
Docs: ni-sh-a-char.github.io/RESENTMENT---kernel/

#operatingsystems #kernel #systemsprogramming #riscv #arm #security #opensource"""

REDDIT_TITLE = "RESENTMENT 2.0.0 — a capability-secure kernel for x86_64, ARM64 and RISC-V, where every capability carries an expiry inside its MAC"

REDDIT_BODY = """Two years ago this was a hobby kernel that printed one line at boot. I've
just tagged v2.0.0 and it's a different project.

**What it is**

A from-scratch kernel for x86_64, aarch64 and riscv64. Not a Unix clone. It's
built around three properties a conventional kernel can't really express:

1. **Authority expires.** No ambient authority anywhere. Every capability is
   checked four ways on use — type, rights, generation, and a cryptographic
   seal — and the seal carries a validity window *inside the MAC*. Widening it
   by editing the struct invalidates the seal. Derivation only ever weakens:
   rights are intersected with the parent's and the lifetime clamped to what
   the parent has left.

2. **The system is a Merkle DAG.** Every kernel object is a node with a
   SHA-256 digest over a canonical encoding of its fields plus the sorted
   digests of its children. Timestamps, pointers and IDs are excluded on
   purpose, so two machines in the same configuration hash identically. You
   get diffing, attestation and deterministic replay from that one property.

3. **Inference is a scheduling class.** A token loop never sleeps for a human,
   so sleep-credit heuristics give it nothing, and it has a visible rate, so
   batch is wrong too. `SCHED_INFERENCE` admits it with a declared rate, gives
   it a per-period budget, and demotes rather than drops it on overrun.

**It runs a transformer**

Not a wrapper around one - the forward pass itself, through the kernel's own
operators: embed, RMSNorm, QKV projections, RoPE, attention over the KV
history, gated FFN, logits, argmax. On a thread the deadline admission
controller accepted.

    resentment> .infer 24
      generated      24 tokens in 130 ms
      rate           193 tokens/sec
      class          SCHED_INFERENCE, 50 Hz declared, 5000 us budget, admitted
      deadline miss  0

The fixture model has pseudo-random weights and the tokens are meaningless by
construction. The claim is about the path, not the output.

**Status, honestly**

Boots to an interactive shell on all three architectures. SMP on all three
(ACPI MADT + a real-mode AP trampoline; PSCI `CPU_ON`; SBI HSM `hart_start`),
tested on four and eight cores. Ring 3 with an ELF64 loader works on x86_64
only — ARM64 and RISC-V need MMU work first and the docs say exactly what's
left. No PCI, no network stack.

**Verification**

    make test           1440 assertions against the real sources on the host
    make qemu-test-all  6 targets (3 arches x single-core and -smp 4)
    make kaalka-check   crypto byte-identical to the reference implementation

Plus seven self-tests on every boot, on the machine about to be trusted.

**The bugs were the best part**

- The scheduler put a thread back on the run queue *before* its stack pointer
  was saved. Two cores could run one stack. Presented as a page fault on cpu1
  in unrelated code, minutes later.
- The x86_64 syscall entry destroyed rdi/rsi/rdx/r8/r9/r10 — it pushed them to
  build its argument block and then dropped it with `add rsp, 7*8`. SYSCALL
  only clobbers rcx and r11, so callers keep live values there.
- The CSPRNG ran on an all-zero key: entropy went into a pool that was only
  folded into the key once an estimate crossed 128 bits, which never happens
  without a hardware RNG.
- On RISC-V, GNU ld relaxes `la rd, sym` to `addi rd, gp, off`. I had code
  before `gp` was set. Every hart computed garbage and parked. Invisible under
  LLD, fatal on hardware.
- **`sched_tick()` was never called on ARM64.** It was driven by a hardcoded
  test for interrupt line zero, which is the x86 timer and nothing else's.
  ARM's is a PPI on line 27. That port had no preemption, no slice accounting,
  and `sched_sleep_ms` never returned. Nothing in the suite slept, so it went
  unnoticed for the life of the port.

Apache 2.0. No dependencies — `make toolchain` fetches a portable zig+nasm and
builds all three architectures, or it uses your system compiler if you have
one.

Code: https://github.com/ni-sh-a-char/RESENTMENT---kernel
Docs: https://ni-sh-a-char.github.io/RESENTMENT---kernel/

Happy to answer anything, especially about the capability model — that's the
part I'd most like criticised."""

HN_TITLE = "RESENTMENT: a capability-secure kernel where authority expires by construction"

HN_COMMENT = """Author here.

The design bet is that revocation should be a background property rather than
an operation. Every capability carries a MAC over its own fields plus a
validity window, so a permission nobody remembered to revoke stops working on
its own, and editing the struct to widen the window invalidates the seal.

The part I'd most like picked apart: the seal's security rests on the platform
clock. An attacker who controls the clock can widen a window. I've written that
down in SECURITY.md rather than hoping nobody notices, but I don't have a good
answer beyond "use a monotonic source the attacker doesn't control", and I'd
like one.

Two other things I'd rather say than have found:

The clock-angle cipher the temporal layer is named after supplies the
*schedule* — epoch keys, seal windows, replay defence. Confidentiality and
integrity come from ChaCha20 and HMAC-SHA256. The clock-angle stream alone is
an additive cipher over a small key space and is not used to protect kernel
objects.

Ring 3 works on all three architectures, and getting there was mostly a lesson
in compilers not honouring their own flags. On ARM64, `-mgeneral-regs-only` is
supposed to keep SIMD out of the kernel; clang emits it in memcpy anyway, and a
page fault on a NEON pair store lost exactly 32 bytes because the exception path
doesn't save NEON - reasonably, since nothing outside kernel/ai should have it.

What's genuinely missing: no PCI, no block driver, no network stack. A process
can't spawn another one yet. And arch_pgtable_destroy leaks every page-table
level below the root, which is fine while processes are rare and won't be.

https://github.com/ni-sh-a-char/RESENTMENT---kernel"""

SHORT_BLURB = """RESENTMENT 2.0.0 — a capability-secure, AI-native kernel.

Authority expires by construction. The whole machine has one SHA-256 root
digest. Inference is a scheduling class.

x86_64 + ARM64 + RISC-V, SMP on all three, 1440 + 180 tests, zero dependencies.

github.com/ni-sh-a-char/RESENTMENT---kernel"""

ONE_LINER = ("A capability-secure, AI-native kernel for x86_64, ARM64 and "
             "RISC-V — where every permission carries its own expiry date.")

INSTAGRAM = """Two years ago: a hobby kernel that printed one line at boot.
Today: v2.0.0.

RESENTMENT is a capability-secure, AI-native kernel.

→ Authority expires. Every permission carries a cryptographic seal with its
  deadline inside. Forget to revoke it and it stops working anyway.
→ The whole machine has one hash. Two machines in the same state produce the
  same digest.
→ Inference is a scheduling class, not an afterthought.

Runs on x86_64, ARM64 and RISC-V. Uses every core. 1,608 automated checks.
Zero dependencies.

Built in the open, Apache 2.0, link in bio.

#osdev #kernel #systemsprogramming #riscv #arm64 #cybersecurity #opensource
#lowlevelprogramming #computerscience #buildinpublic"""


# ------------------------------------------------------------------ build

def build():
    os.makedirs(OUT, exist_ok=True)
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = BODY
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)

    for s in doc.sections:
        s.top_margin = Inches(0.85)
        s.bottom_margin = Inches(0.85)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # ---------------------------------------------------------- cover
    para(doc, "RESENTMENT " + VERSION, size=30, bold=True, space_after=2)
    para(doc, "Social kit — ready-to-post copy, assets and talking points",
         size=12, colour=MUTED, space_after=14)

    cover = os.path.join(OUT, "social-preview.png")
    if os.path.exists(cover):
        doc.add_picture(cover, width=Inches(6.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(doc, "media/social-preview.png — 1280x640", size=8.5,
             colour=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    para(doc, ONE_LINER, size=12.5, italic=True, colour=INK, space_after=12)

    kv_table(doc, [
        ("Repository",    REPO),
        ("Documentation", DOCS),
        ("Licence",       "Apache 2.0"),
        ("Architectures", "x86_64, aarch64, riscv64 — all boot, all SMP"),
        ("Verification",  "1440 host assertions + 180 through the shell under "
                          "QEMU across 6 targets"),
        ("Dependencies",  "none — builds from its own sources plus a compiler"),
        ("Tags",          "v1.0.0 (the original hobby kernel) and v2.0.0 (this)"),
    ])

    doc.add_page_break()

    # ------------------------------------------------------ how to use
    h1(doc, "How to use this document")
    para(doc, "Everything in a shaded box is meant to be copied out verbatim. "
              "Everything outside one is context for you, not for the post.")
    bullet(doc, "The numbers are current as of " + VERSION + ". If you post "
                "later and the suite has grown, regenerate this with "
                "tools/mkdocx.py rather than editing by hand.")
    bullet(doc, "Nothing here overclaims. The limitations section is included "
                "on purpose — the fastest way to lose a technical audience is "
                "to be caught overstating, and the fastest way to win one is "
                "to name your own gaps first.")
    bullet(doc, "Pick one platform to lead with. Posting the same thing "
                "everywhere within the hour reads as broadcast; spacing it "
                "over a few days reads as a person.")

    h2(doc, "Which image goes where")
    kv_table(doc, [
        ("social-preview.png",  "1280x640 — GitHub repository social preview. "
                                "Settings > General > Social preview."),
        ("social-wide.png",     "1600x900 — X, YouTube, anything 16:9. Carries "
                                "a real boot transcript."),
        ("social-linkedin.png", "1200x627 — LinkedIn and Facebook link cards."),
    ])
    para(doc, "All three are generated by tools/mksocial.py and rebuild in "
              "about a second. The fonts are Inter and JetBrains Mono, both "
              "SIL Open Font Licence, which is what makes it legitimate to "
              "bake them into a distributable image.", colour=MUTED, size=10)

    doc.add_page_break()

    # ------------------------------------------------------------- X
    h1(doc, "X / Twitter")
    h3(doc, "Single post")
    para(doc, "Lead with this if you only post once. Attach social-wide.png.",
         colour=MUTED, size=10)
    block(doc, X_MAIN)

    h3(doc, "Thread")
    para(doc, "Eleven posts. The bug posts (6 to 9) are the ones that travel — "
              "specific, verifiable, and they cost you nothing to admit.",
         colour=MUTED, size=10)
    block(doc, X_THREAD)

    doc.add_page_break()

    # ------------------------------------------------------ LinkedIn
    h1(doc, "LinkedIn")
    para(doc, "Longer, and written for people who will not read the code. The "
              "bold runs are Unicode mathematical bold, because LinkedIn has "
              "no formatting — paste it as-is and it renders.",
         colour=MUTED, size=10)
    block(doc, LINKEDIN, mono=False)

    doc.add_page_break()

    # -------------------------------------------------------- Reddit
    h1(doc, "Reddit")
    para(doc, "r/osdev is the right first home. r/rust, r/programming and "
              "r/RISCV are secondary; r/programming will be harsher and wants "
              "the bug list up front.", colour=MUTED, size=10)
    h3(doc, "Title")
    block(doc, REDDIT_TITLE)
    h3(doc, "Body (markdown)")
    block(doc, REDDIT_BODY)

    doc.add_page_break()

    # --------------------------------------------------- Hacker News
    h1(doc, "Hacker News")
    h3(doc, "Title")
    para(doc, "Under 80 characters, no version number, no exclamation mark.",
         colour=MUTED, size=10)
    block(doc, HN_TITLE)

    h3(doc, "First comment, posted by you immediately after submitting")
    para(doc, "This is the single highest-leverage thing in the document. HN "
              "rewards an author who names the weakest part of their own work "
              "before anybody else does.", colour=MUTED, size=10)
    block(doc, HN_COMMENT)

    doc.add_page_break()

    # ---------------------------------------------------- short form
    h1(doc, "Short form")
    h3(doc, "Discord, Slack, Mastodon, anywhere with a small box")
    block(doc, SHORT_BLURB)
    h3(doc, "One line, for a bio or a talk abstract")
    block(doc, ONE_LINER)
    h3(doc, "Instagram / Threads")
    block(doc, INSTAGRAM, mono=False)

    doc.add_page_break()

    # ------------------------------------------------- talking points
    h1(doc, "Talking points")
    para(doc, "For replies and comments. Each is a claim you can defend with "
              "something in the repository.", colour=MUTED, size=10)

    h2(doc, "Why not just use Linux?")
    para(doc, "You should, for almost everything. RESENTMENT exists to explore "
              "three properties Linux cannot express without being a different "
              "kernel: authority that expires by construction, a whole-system "
              "state hash that is a function of meaning rather than layout, "
              "and inference as a first-class scheduling class. If those are "
              "not properties you need, Linux is a better answer.")

    h2(doc, "What does “capability-secure” actually buy?")
    para(doc, "A process here holds no ambient authority at all. The demo "
              "program in ring 3 can compute, print and exit — it cannot open "
              "a file, because opening a file means presenting a capability "
              "for a directory and it holds none. The interesting question "
              "stops being “what can an attacker reach” and becomes "
              "“what was this process handed”.")

    h2(doc, "Is the clock-angle cipher doing the real work?")
    para(doc, "No, and the project says so in its own documentation. Kaalka "
              "supplies the schedule: epoch keys, seal windows, replay "
              "defence. Confidentiality and integrity come from ChaCha20 and "
              "HMAC-SHA256. The clock-angle stream alone is an additive cipher "
              "over a small key space and is not used to protect kernel "
              "objects. Saying that plainly is part of the design.")

    h2(doc, "How is it tested, really?")
    para(doc, "1440 assertions compile the real kernel sources for the host "
              "against a synthetic machine — they found a live alignment bug "
              "in the slab allocator on their first run. Then six QEMU "
              "targets: each architecture booted single-core and again with "
              "four cores, driven through its own shell over a serial socket. "
              "Then seven self-tests on every boot, on the machine about to be "
              "trusted, because a kernel whose crypto is wrong should not be "
              "able to look healthy.")

    h2(doc, "What is genuinely not done")
    bullet(doc, "No SYS_TASK_SPAWN: a process cannot start another one, "
                "only the kernel and the shell can.")
    bullet(doc, "No PCI enumeration, no block driver, no network stack. The "
                "device model exists; the drivers do not.")
    bullet(doc, "The AI subsystem parses GGUF and has the tensor, operator, "
                "accelerator and KV-cache layers, but does not yet run a model "
                "end to end.")
    bullet(doc, "Temporal seals depend on the platform clock. An attacker who "
                "controls it can widen a validity window.")
    bullet(doc, "One shared run queue rather than per-CPU queues. Deliberate, "
                "and documented as a trade rather than an oversight.")

    doc.add_page_break()

    # ---------------------------------------------------------- facts
    h1(doc, "Fact sheet")
    para(doc, "Every number here is reproducible with a command in the "
              "repository.", colour=MUTED, size=10)
    kv_table(doc, [
        ("Architectures",      "x86_64, aarch64, riscv64"),
        ("Boots to a shell",   "all three, roughly 120 ms"),
        ("SMP",                "all three; tested on 4 and 8 cores"),
        ("Wall clock",         "all three — CMOS, PL031, goldfish"),
        ("Ring 3 userspace",   "all three - ELF64 loader, per-process address spaces, syscalls"),
        ("Inference",          "all three - a real transformer forward pass, admitted by deadline admission control"),
        ("Scheduling classes", "4 — realtime, inference, interactive, batch"),
        ("Capability types",   "24, each sealed and time-bounded"),
        ("SHE builtins",       "29, behind 11 permission grants"),
        ("Host assertions",    "1440 (make test)"),
        ("QEMU assertions",    "180 across 6 targets (make qemu-test-all)"),
        ("Boot self-tests",    "7, on every boot"),
        ("Kaalka cross-check", "100% byte-identical to the reference"),
        ("Dependencies",       "zero"),
        ("Lines of C",         "about 26,000"),
        ("Licence",            "Apache 2.0"),
    ], widths=(2.3, 4.2))

    h2(doc, "The three projects it is built out of")
    kv_table(doc, [
        ("Kaalka",    "Time-driven cryptography from clock-hand angles. Becomes "
                      "the temporal authority layer: every capability seal, IPC "
                      "envelope and snapshot is bound to a window and an epoch key."),
        ("WebWeaveX", "Deterministic runtime graphs with stable node identity. "
                      "Becomes the kernel's own state model."),
        ("SHE",       "A language whose programs start with zero permissions. "
                      "Becomes the system shell, sandboxed by the kernel's "
                      "capability space rather than by an interpreter."),
    ], widths=(1.3, 5.2))

    h2(doc, "Posting checklist")
    for item in [
        "Set the GitHub social preview (Settings > General > Social preview) "
        "to media/social-preview.png before posting anywhere.",
        "Check the docs site is live: " + DOCS,
        "Have the repository on a screen — the first question is always "
        "“show me the code that does X”.",
        "Reply to every substantive comment in the first two hours. After "
        "that the thread is cold.",
        "When someone finds a real problem, say so and link the issue. That "
        "single behaviour converts more sceptics than any feature.",
    ]:
        bullet(doc, item)

    path = os.path.join(OUT, "RESENTMENT-%s-social-kit.docx" % VERSION)
    doc.save(path)
    print("  %s  (%.1f KiB)" % (os.path.relpath(path, ROOT),
                                os.path.getsize(path) / 1024.0))
    return 0


if __name__ == "__main__":
    sys.exit(build())
